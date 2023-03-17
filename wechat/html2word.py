from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import os
from urllib.request import urlretrieve
from docx import Document
from docx.shared import Inches
def trimName(name):
    return name.replace(' ', '').replace('|', '，').replace('\\', '，').replace('/', '，').replace(':', '，').replace('*', '，').replace('?', '，').replace('<', '，').replace('>', '，').replace('"', '，').replace('\n', '，').replace('\r', '，').replace(',', '，').replace('\u200b', '，').replace('\u355b', '，').replace('\u0488', '，').replace('•','')
def createDoc(title,publish_time,doc,doc_name):
    document = Document()
    #添加标题，并设置级别，范围：0 至 9，默认为1
    document.add_heading(title, 0)
    #添加段落，文本可以包含制表符（\t）、换行符（\n）或回车符（\r）等
    document.add_paragraph(publish_time)
    for i in range(len(doc)):
        if doc[i][1]=='' or doc[i][2]==0:
            continue
        if doc[i][0]=='title':
            document.add_heading(doc[i][1], level=1)
        elif doc[i][0]=='content':
            document.add_paragraph(doc[i][1])
        elif doc[i][0]=='image':
            document.add_picture(doc[i][1],width=Inches(5.9))
        else:
            raise
    document.save('word/'+doc_name)
    print("保存成功",doc_name)
def convert(url):
    opt = Options()
    opt.add_argument('--no-sandbox')                # 解决DevToolsActivePort文件不存在的报错
    opt.add_argument('--disable-gpu')               # 谷歌文档提到需要加上这个属性来规避bug
    opt.add_argument('blink-settings=imagesEnabled=false')      # 不加载图片，提升运行速度
    opt.add_argument('--headless')                  # 浏览器不提供可视化界面。Linux下如果系统不支持可视化不加这条会启动失败
    brower = webdriver.Chrome(options=opt)
    brower.get(url)
    doc = []
    title = brower.find_element_by_id('activity-name') #纯文字没有标题
    publish_time = brower.find_element_by_id('publish_time')
    publish_time = publish_time.text
    date = publish_time[:10]
    doc_name = trimName(date+'-'+title.text)
    images_folder_name = doc_name
    subtitles = brower.find_elements_by_tag_name('strong')
    contents = brower.find_elements_by_tag_name('p')
    images = brower.find_elements_by_class_name("wxw-img")
    if not os.path.exists(f'{images_folder_name}'):
        os.mkdir(f'./{images_folder_name}')
    # for subtitle in subtitles:
    #     doc.append(['title',subtitle.text,subtitle.location['y']])
    for content in contents:
        chiildContents = content.find_elements_by_tag_name('p')
        chiildContents2 = content.find_elements_by_tag_name('em')
        if not chiildContents and not chiildContents2:
            doc.append(['content',content.text,content.location['y']])
        
    for id,image in enumerate(images):
        img_url = image.get_attribute('data-src')
        img_name = f'./{images_folder_name}/{id}.png'
        urlretrieve(img_url, img_name)
        if image.size['height']>50 and image.size['width']>50:
            doc.append(['image',img_name,image.location['y']])
    doc.sort(key=lambda x:x[2])
    # print(doc,title.text,publish_time)
    doc_name += ".docx"
    createDoc(title.text,publish_time,doc,doc_name)
urls = [] 
filename = input("请输入文件名：")
with open(f'{filename}', encoding='utf-8') as f:
     contents = f.read()
urls = contents.split("\n")
# f = open(f'{filename}', encoding='utf-8')
# csv_reader = csv.reader(f)
print(len(urls))
# urls=['https://mp.weixin.qq.com/s/BTrQnsZ6eKWXf5n1mWhv7A']
if not os.path.exists(f'word'):
    os.mkdir(f'./word')
for i in urls:
	print("开始转换",i)
	try:
		convert(i)
	except Exception as e:
		print("转换失败",e,i)
# for line in csv_reader:
#     try:
#         convert(line[2])
#     except Exception as e:
#         print("转换失败",e,line[2])