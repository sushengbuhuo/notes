import asyncio,os,requests
from pyppeteer import launch
#/share/88dfb803-15f0-43fc-b6c3-3623d1a2d5a7
async def export_pdf(url, output_path):
    browser = await launch()
    page = await browser.newPage()
    await page.goto(url, {'waitUntil': 'networkidle0'})
    await page.pdf({'path': output_path, 'format': 'A4'})
    
    # 关闭当前页面
    await page.close()

# 批量导出 PDF
async def batch_export_pdfs(urls, output_folder):
    tasks = []
    
    for index, url in enumerate(urls):
        output_path = f'{output_folder}/output_{index + 1}.pdf'
        task = export_pdf(url, output_path)
        tasks.append(task)

    # 并行执行任务
    await asyncio.gather(*tasks)

# 用法示例
urls_to_export = ['https://example1.com', 'https://example2.com', 'https://example3.com']
output_folder = 'output_folder'

# 异步运行批量导出
# asyncio.get_event_loop().run_until_complete(batch_export_pdfs(urls_to_export, output_folder))
from docx import Document
from bs4 import BeautifulSoup

def html_to_docx(html_content, output_path):
    # 创建一个新的 Word 文档
    doc = Document()

    # 使用 BeautifulSoup 解析 HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # 遍历 HTML 中的段落，并将其添加到 Word 文档中
    for paragraph in soup.find_all('p'):
        doc.add_paragraph(paragraph.get_text())

    # 保存 Word 文档
    doc.save(output_path)

# 用法示例
html_content = """
<html>
  <body>
    <p>This is a paragraph in HTML.</p>
    <p>Another paragraph in HTML.</p>
  </body>
</html>
"""

output_path = 'output.docx'

# html_to_docx(html_content, output_path)
from docxtpl import DocxTemplate
from bs4 import BeautifulSoup

def html_to_docx_with_images(html_content, output_path):
    # 创建一个新的 Word 文档模板
    doc = DocxTemplate("template.docx")  # 你可以替换为自己的 Word 模板

    # 使用 BeautifulSoup 解析 HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # 遍历 HTML 中的段落，并将其添加到 Word 文档模板中
    for paragraph in soup.find_all('p'):
        text = paragraph.get_text()
        doc.render({'paragraph': text})
        doc.save(output_path)

# 用法示例
html_content_with_images = """
<html>
  <body>
    <p>This is a paragraph in HTML.</p>
    <p>Another paragraph in HTML.</p>
    <img src="path/to/image.jpg" alt="Image Alt Text">
  </body>
</html>
"""

output_path_with_images = 'output_with_images.docx'

# html_to_docx_with_images(html_content_with_images, output_path_with_images)
import pyperclip
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

def copy_html_to_word(html_content, output_path):
    # 将 HTML 内容复制到剪贴板
    pyperclip.copy(html_content)

    # 创建一个新的 Word 文档
    doc = Document()

    # 粘贴剪贴板中的内容到 Word 文档
    doc.add_paragraph(pyperclip.paste())

    # 保存 Word 文档
    doc.save(output_path)

# 用法示例
html_content = """
<html>
  <body>
    <p>This is a paragraph in HTML.</p>
    <p>Another paragraph in HTML.</p>
    <img src="" alt="Image Alt Text">
  </body>
</html>
"""

output_path = 'pyoutput.docx'
if not os.path.exists('images'):
    os.mkdir('images')
# copy_html_to_word(html_content, output_path)
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from docx import Document
import time
headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/53.0.2785.116 Safari/537.36 QBCore/4.0.1301.400 QQBrowser/9.0.2524.400 Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/53.0.2875.116 Safari/537.36 NetType/WIFI MicroMessenger/7.0.5 WindowsWechat"
    }
# https://googlechromelabs.github.io/chrome-for-testing/  https://registry.npmmirror.com/binary.html?path=chrome-for-testing http://chromedriver.storage.googleapis.com/index.html chrome://version/
# /c/5be397b3-393d-46f4-ba18-d43aa7764b1d
def copy_html_to_word(html_url, output_path):
    # 配置 Chrome 选项，使浏览器在无头模式下运行
    chrome_options = Options()
    chrome_options.add_argument("--headless")

    # 初始化 WebDriver
    driver = webdriver.Chrome(options=chrome_options)

    # 打开 HTML 页面
    driver.get(html_url)

    # 等待一段时间以确保页面加载完成
    time.sleep(3)

    # 获取页面的文本内容
    text_content = driver.find_element_by_tag_name('body').text
    print(text_content)
    # 获取页面的图片链接列表
    img_elements = driver.find_elements_by_tag_name('img')
    img_links = [img.get_attribute('data-src') for img in img_elements]
    img_links = list(filter(lambda x: x is not None, img_links))
    img_links = [x for x in img_links if x is not None]
    print(img_links)
    # 创建 Word 文档
    doc = Document()

    # 添加文本内容
    doc.add_paragraph(text_content)
    num = 0
    # 添加图片
    for img_link in img_links:
        num+=1
        img_data = requests.get(img_link,headers=headers)
        with open('images/'+str(num)+'.jpg','wb') as f6:
            f6.write(img_data.content)
        doc.add_picture('images/'+str(num)+'.jpg', width=300)  # 调整图片宽度适应页面

    # 保存 Word 文档
    doc.save(output_path)

    # 关闭浏览器
    driver.quit()

# 用法示例
html_url = 'https://'
output_path = 'wordoutput.docx'
# copy_html_to_word(html_url, output_path)
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from io import BytesIO
from PIL import Image
# 输入你想要复制的网页地址
url = "https://"

# 发送HTTP请求获取网页内容
response = requests.get(url, headers=headers)
html_content = response.content

# 使用Beautiful Soup解析网页内容
soup = BeautifulSoup(html_content, 'html.parser')

# 创建一个Word文档
doc = Document()

# 提取网页文字内容
text_content = soup.get_text()
doc.add_paragraph(text_content)

# 提取网页图片并插入到Word文档中
img_tags = soup.find_all('img')
for img_tag in img_tags:
    img_url = img_tag['data-src']
    img_response = requests.get(img_url)
    img_data = BytesIO(img_response.content)
    img = Image.open(img_data)

    # 设置图片大小（可根据需要调整）
    width, height = img.size
    max_width = 400
    if width > max_width:
        ratio = max_width / width
        width = max_width
        height = int(height * ratio)

    # 将图片插入到Word文档中
    doc.add_picture(img_data, width=width, height=height)

# 保存Word文档
doc.save('pyoutput.docx')


