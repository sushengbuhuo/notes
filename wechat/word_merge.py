import os
from docx import Document
from docxcompose.composer import Composer
import sys
def merge_doc(source_file_path_list, target_file_path):
    page_break_doc = Document()
    page_break_doc.add_page_break()
    target_doc = Document(source_file_path_list[0])
    target_composer = Composer(target_doc)
    for i in range(len(source_file_path_list)):
        # 跳过第一个
        if i == 0:
            continue
        # 填充分页符
        print(source_file_path_list)
        try:
            target_composer.append(page_break_doc)
        except Exception as err:
            print(err)
            return
        # 拼接
        f = source_file_path_list[i]
        target_composer.append(Document(f))
    # 保存
    target_composer.save(target_file_path)

if __name__ == '__main__':
    current_path = os.path.abspath('.')
    source_path = rf'{current_path}'
    target_file = rf'{current_path}/文档合集.docx'
    source_file_list = os.listdir(source_path)
    new_list = []
    for item in source_file_list:
        if item.endswith('.docx'):
            new_list.append(item)
    # 文件列表
    source_file_list_all = []
    for file in new_list:
        source_file_list_all.append(source_path + '\\' + file)
    # 合并
    try:
        merge_doc(source_file_list_all, target_file)
    except Exception as e:
        print(e)
