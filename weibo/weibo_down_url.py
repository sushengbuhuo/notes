import requests,re,os,time,html,sys,csv
import random
import traceback,urllib3
from docx import Document
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
ALPHABET = "0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ"
def base62_encode(num, alphabet=ALPHABET):
    num = int(num)
    if num == 0:
        return alphabet[0]
    arr = []
    base = len(alphabet)
    while num:
        rem = num % base
        num = num // base
        arr.append(alphabet[rem])
    arr.reverse()
    return ''.join(arr)

def base62_decode(string, alphabet=ALPHABET):
    string = str(string)
    num = 0
    idx = 0
    for char in string:
        power = (len(string) - (idx + 1))
        num += alphabet.index(char) * (len(alphabet) ** power)
        idx += 1

    return num
def reverse_cut_to_length(content, code_func, cut_num=4, fill_num=7):
    content = str(content)
    cut_list = [content[i - cut_num if i >= cut_num else 0:i] for i in range(len(content), 0, (-1 * cut_num))]
    cut_list.reverse()
    result = []
    for i, item in enumerate(cut_list):
        s = str(code_func(item))
        if i > 0 and len(s) < fill_num:
            s = (fill_num - len(s)) * '0' + s
        result.append(s)
    return ''.join(result)
def get_cookie():
    cookie = ''
    if os.path.exists('cookie.txt'):
        with open('cookie.txt', encoding='utf-8') as f:
            cookie = f.read().replace('\n','')
    return cookie
cookie = get_cookie()
headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/53.0.2785.116 Safari/537.36 QBCore/4.0.1301.400 QQBrowser/9.0.2524.400 Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/53.0.2875.116 Safari/537.36 NetType/WIFI MicroMessenger/7.0.5 WindowsWechat",
        'referer': 'https://weibo.com/1744395855/NkD5bjvPC',
        "Cookie":cookie,
    }
# url = input('请输入微博链接：')
# res = requests.get(html.unescape(url),proxies={'http': None,'https': None},verify=False, headers=headers)
# title = re.search(r'<title>(.*?)</title>',res.text).group(1)
# title = re.search(r'<div class="Feed_body_3R0rO">(.*?)<!----></div',res.text).group(1)
# weibo_time = re.search(r'<span class="time".*>(.*?)</span>',res.text).group(1)
# if not weibo_time.startswith('20'):
# 	weibo_time=time.strftime('%Y')+'-'+weibo_time.strip().split(' ')[0]
# content = '<!DOCTYPE html><html><head><meta charset="utf-8"></head><body><h1>%s</h1><h3>%s</h3>%s</body></html>' % (
#             res2['post']['post_title'], '发布日期：'+res2['post']['post_publish_time']+'   原文链接：https://emcreative.eastmoney.com/app_fortune/article/index.html?postId=')
# try:
# 	with open('html/'+res2['post']['post_publish_time'][0:10]+'-'+trimName(res2['post']['post_title'])+'.html', 'w', encoding='utf-8') as f:
# 		f.write(content)
# except Exception as err:
# 	with open('html/'+str(random.randint(100,1000))+'.html', 'w', encoding='utf-8') as f:
# 		f.write(content)
# import asyncio,os
# from pyppeteer import launch
# import tkinter,time
from datetime import datetime
# from playwright.sync_api import sync_playwright
def replace_invalid_chars(filename):
    invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*','\n','#']
    for char in invalid_chars:
        filename = filename.replace(char, ' ')
    return filename
def get_history():
    history = []
    with open('weibo_history.txt', 'a+') as f:
        f.seek(0)
        lines = f.readlines()
        for line in lines:
            history.append(line.strip())
    return history

def save_history(url):
    with open('weibo_history.txt', 'a+') as f:
        f.write(url.strip() + '\n')
def main():
    f = open(f'微博.csv', encoding='UTF8')
    csv_reader = csv.reader(f)
    # for root, dirs, files in os.walk('.'):
    num = 0
    history = get_history()
    # browser = await launch(headless=False)
    # browser = await launch()# {'args': ['--disable-infobars'],'userDataDir': './userdata'} 登录后保存cookie
    for line in csv_reader:
        if line[0] in history:
            print('已经下载过:'+line[0])
            continue
        time.sleep(random.randint(2, 6))
        # if num>10:
        #     break
        if '微博' in line[0]:
            continue
        num +=1
        try:
            m=re.search(r'https://www\.weibo\.com/\d+/(.*)',line[0]).group(1)
            # mid=reverse_cut_to_length(m, base62_decode, 4, 7)
            # url2 =f'https://m.weibo.cn/detail/{mid}'
            # page = await browser.newPage() time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()
            # page.setDefaultNavigationTimeout(60000)  # 设置为60秒
            url=f'https://weibo.com/ajax/statuses/show?id={m}&locale=zh-CN'#https://www.weibo.com/ajax/statuses/extend?id=5049927780796644
            res = requests.get(html.unescape(url),proxies={'http': None,'https': None},verify=False, headers=headers).json()
            # dt_obj = datetime.strptime(res['created_at'], '%a %b %d %H:%M:%S %z %Y')
            # created_at = dt_obj.strftime('%Y-%m-%d %H:%M:%S')
            created_at=line[7]
            dt_obj = datetime.strptime(created_at, '%Y-%m-%d %H:%M')
            date =  dt_obj.strftime('%m月%d日');year=created_at[0:4];minute=created_at[11:16].replace(':','：')
            if not os.path.exists(year):
                os.mkdir(year)
            if not os.path.exists(f'{year}/{date}'):
                os.mkdir(f'{year}/{date}')
            if not os.path.exists(f'{year}/{date}/{minute}'):
                os.mkdir(f'{year}/{date}/{minute}')
            with open(f'{year}/{date}/{minute}/{m}.txt', 'a+', encoding='utf-8') as f2:
                f2.write(res['text_raw'])
            document = Document()
            # document.add_heading('文本内容', 0)
            document.add_paragraph(res['text_raw'])
            document.save(f'{year}/{date}/{minute}/{m}.docx')
            print('开始下载',line[0],created_at)
            print(res['text_raw'])
            if res['pic_num'] > 0:
                # if not os.path.exists(f'{date}/image'):
                #     os.mkdir(f'{date}/image')
                for j,k in res['pic_infos'].items():
                    print('图片:',k['largest']['url'])
                    img_data = requests.get(k['largest']['url'].replace('/large/','/oslarge/'),headers=headers,timeout=5)
                    with open(f'{year}/{date}/{minute}/'+j+'.jpg','wb') as f3:
                        f3.write(img_data.content)
            if 'page_info' in res and 'media_info' in res.get('page_info') and 'playback_list' in res.get('page_info').get('media_info'):
                # if not os.path.exists(f'{date}/video'):
                    # os.mkdir(f'{date}/video')
                video_url = res.get('page_info').get('media_info').get('playback_list')[0]['play_info']['url']
                title=res.get('page_info').get('media_info').get('name')+res.get('page_info').get('object_id')
                print('视频:',video_url)
                video_data = requests.get(video_url,headers=headers,verify=False,timeout=10)
                with open(f'{year}/{date}/{minute}/'+replace_invalid_chars(title)+'.mp4','wb') as f4:
                    f4.write(video_data.content)
            save_history(line[0])
            # url = "https://mp.weixin.qq.com/s/S24LAiMtAfdGS9XM0ZMU2A"
            # 查看当前 桌面视图大小 https://miyakogi.github.io/pyppeteer/reference.html
            #tk = tkinter.Tk()
            #width = tk.winfo_screenwidth()
            #height = tk.winfo_screenheight()
            #tk.quit()
            #print(f'设置窗口为：width：{width} height：{height}')
            # await asyncio.sleep(4.12)
            # 设置网页 视图大小
            #await page.setViewport(viewport={'width': width, 'height': height})
            #参数 https://ld246.com/article/1566221786951  https://blog.csdn.net/weixin_45961774/article/details/112848584
            #await page.evaluateOnNewDocument('function(){Object.defineProperty(navigator, "webdriver", {get: () => undefined})}')
            # await page.setUserAgent('Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/71.0.3542.0 Safari/537.36')
            # cookie_dict = {}
            # for pair in cookie_pairs:
            # 	value = pair.split("=")
            # 	cookie_dict[value[0]] = value[1]
            # print(cookie_dict)
            # await page.setCookie(cookie_dict)
            # print(url)
            # await page.goto('https://m.weibo.cn/detail/5000660202553386')
            # page_text = await page.content()  # 页面内容
            # cookies = await page.cookies()
            # 等待页面加载完成
            # await page.waitForNavigation()
            # await page.pdf({"path": 'pdf/'+line[7][0:10]+'_'+line[1][0:20]+'.pdf', "format": 'A4'})
            # await page.close()
        except Exception as e:
            pass
            # if not os.path.exists('failed'):
            #     os.mkdir('failed')
            # shutil.copy(name, 'failed')
            # Navigation Timeout Exceeded: 30000 ms exceeded
            # print('下载失败',e,line[0])#;raise Exception("抓取失败了："+line[0])
    # break
    # htmls += [name for name in files if name.endswith(".html")]
    # await browser.close()
# async def main2():
#     for root, dirs, files in os.walk('.'):
#         files.sort(reverse = True)
#         for name in files:
#             if name.endswith(".html"):
#                 print(name,time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()))
#                 try:
#                     pdf_file = 'pdf/'+name.replace('.html', '')+'.pdf'
#                     os.system(f'playwright pdf {name} {pdf_file}')
#                     # os.system(f'wkhtmltopdf {name} {pdf_file}')
#                 except Exception as e:
#                     print(e)
#         # break
#         # htmls += [name for name in files if name.endswith(".html")]
#     # print(htmls)
main()
    